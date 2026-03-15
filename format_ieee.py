"""
format_ieee.py — Enforce strict IEEE conference/journal formatting on H5-OmniFusion paper
=====================================================================================
Applies IEEE formatting rules to H5_OMNIFUSION_JOURNAL_PAPER_V4.docx:

IEEE Formatting Rules (Transactions/Journals style):
  - Page: US Letter (8.5 x 11 in)
  - Margins: Top = 0.75in (19mm), Bottom = 1.0in (25.4mm),
             Left = Right = 0.625in (15.875mm) [first page bottom often 1in]
  - Two-column layout (except title block): column width 3.5in, gutter 0.25in
  - Title: 24pt, centered, Times New Roman
  - Author names: 11pt, centered
  - Affiliations: 10pt, centered, italic
  - Abstract: 9pt bold "Abstract—" followed by 9pt text, justified, single column
  - Keywords: 9pt bold italic "Keywords—" followed by 9pt italic text
  - Section headings (I, II...): 10pt, small caps, centered
  - Subsection headings (A, B...): 10pt, italic, left-aligned
  - Body text: 10pt, Times New Roman, justified, single line spacing
  - Figure captions: 8pt, "Fig." prefix
  - Table captions: 8pt, "TABLE" prefix (above table)
  - Table text: 8pt
  - References: 8pt, Times New Roman, left-aligned
  - Footnotes: 8pt
  - Line spacing: single (exact 12pt for body)
  - First paragraph after heading: no indent; subsequent paragraphs: 0.25in indent

Usage:
    python format_ieee.py

Input:  docs/ml_pipeline/H5_OMNIFUSION_JOURNAL_PAPER_V4.docx
Output: docs/ml_pipeline/H5_OMNIFUSION_JOURNAL_PAPER_V4.docx (overwritten in-place)
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree


# ─────────────────────────────────────────────────────────────────────
# IEEE Constants
# ─────────────────────────────────────────────────────────────────────

IEEE_FONT = "Times New Roman"
IEEE_BODY_SIZE = Pt(10)
IEEE_TITLE_SIZE = Pt(24)
IEEE_AUTHOR_SIZE = Pt(11)
IEEE_AFFILIATION_SIZE = Pt(10)
IEEE_ABSTRACT_SIZE = Pt(9)
IEEE_SECTION_HEADING_SIZE = Pt(10)
IEEE_SUBSECTION_HEADING_SIZE = Pt(10)
IEEE_TABLE_SIZE = Pt(8)
IEEE_REFERENCE_SIZE = Pt(8)
IEEE_CAPTION_SIZE = Pt(8)

IEEE_MARGIN_TOP = Inches(0.75)
IEEE_MARGIN_BOTTOM = Inches(1.0)
IEEE_MARGIN_LEFT = Inches(0.625)
IEEE_MARGIN_RIGHT = Inches(0.625)

IEEE_COL_COUNT = 2
IEEE_COL_SPACE = Inches(0.25)

IEEE_LINE_SPACING = Pt(12)  # Single spacing = exactly 12pt for 10pt font
IEEE_PARA_SPACING_BEFORE = Pt(0)
IEEE_PARA_SPACING_AFTER = Pt(0)
IEEE_FIRST_LINE_INDENT = Inches(0.25)


# ─────────────────────────────────────────────────────────────────────
# Helper functions
# ─────────────────────────────────────────────────────────────────────

def set_run_font(run, name=IEEE_FONT, size=None, bold=None, italic=None, color=None, small_caps=None):
    """Set font properties on a run."""
    run.font.name = name
    # Fix East Asian font fallback
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), name)
    
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if small_caps is not None:
        run.font.small_caps = small_caps


def set_paragraph_format(paragraph, alignment=None, space_before=None, space_after=None,
                         line_spacing=None, line_spacing_rule=None, first_line_indent=None, keep_together=None):
    """Set paragraph format properties."""
    pf = paragraph.paragraph_format
    if alignment is not None:
        pf.alignment = alignment
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    if line_spacing_rule is not None:
        pf.line_spacing_rule = line_spacing_rule
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if keep_together is not None:
        pf.keep_together = keep_together


def set_all_runs_font(paragraph, **kwargs):
    """Set font for all runs in a paragraph."""
    for run in paragraph.runs:
        set_run_font(run, **kwargs)


def is_section_heading(text):
    """Check if this is a major section heading (I. II. III. etc.)"""
    t = text.strip()
    # Matches Roman numerals followed by a period, or specific section names
    patterns = [
        r'^[IVX]+\.\s',           # I. INTRODUCTION, II. RELATED etc.
        r'^ABSTRACT',
        r'^REFERENCES',
        r'^ACKNOWLEDGMENT',
    ]
    return any(re.match(p, t, re.IGNORECASE) for p in patterns)


def is_subsection_heading(text):
    """Check if this is a subsection heading (1.1, 2.3, A., B. etc.)"""
    t = text.strip()
    patterns = [
        r'^\d+\.\d+\s',           # 1.1 Problem Statement
        r'^[A-Z]\.\s',            # A. Feature Extraction
    ]
    return any(re.match(p, t) for p in patterns)


def is_reference(text):
    """Check if this is a reference entry [1], [2], etc."""
    t = text.strip()
    return bool(re.match(r'^\[\d+\]', t))


def is_title_paragraph(index, text):
    """Check if this is the title (paragraph 0, large bold text)."""
    return index == 0 and 'OmniFusion' in text


def is_author_table(table):
    """Check if a table is the author info table (checks all cells in first 2 rows)."""
    author_names = ['Pooja', 'Sivaprakash', 'Karthikeyini', 'Nishanth', 'Nithin',
                    'Computer Science', 'Engineering', 'Krishna', 'Institute']
    for row in table.rows[:2]:
        for cell in row.cells:
            cell_text = cell.text.strip()
            if any(name in cell_text for name in author_names):
                return True
    return False


def is_abstract_paragraph(text):
    """Check if this is the abstract paragraph."""
    t = text.strip()
    return t.startswith('Abstract') or t.startswith('ABSTRACT')


def is_keywords_paragraph(text):
    """Check if this is the keywords paragraph."""
    t = text.strip()
    return t.startswith('Keywords') or t.startswith('KEYWORDS') or t.startswith('Index Terms')


# ─────────────────────────────────────────────────────────────────────
# Main formatting engine
# ─────────────────────────────────────────────────────────────────────

def format_ieee(doc_path):
    """Apply IEEE formatting to the document."""
    print(f"\n{'='*70}")
    print(f"  IEEE Formatting Engine")
    print(f"{'='*70}")
    print(f"\n  Input/Output: {doc_path}")
    
    doc = Document(doc_path)
    changes = []
    
    # ─── 1. Page Setup (all sections) ────────────────────────────────
    print("\n  [1/7] Setting page margins and layout...")
    for i, section in enumerate(doc.sections):
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = IEEE_MARGIN_TOP
        section.bottom_margin = IEEE_MARGIN_BOTTOM
        section.left_margin = IEEE_MARGIN_LEFT
        section.right_margin = IEEE_MARGIN_RIGHT
        section.header_distance = Inches(0.3)
        section.footer_distance = Inches(0.3)
        
        # Set two-column layout for section with body text (section index >= 1)
        sect_pr = section._sectPr
        cols = sect_pr.find(qn('w:cols'))
        if cols is None:
            cols = parse_xml(f'<w:cols {nsdecls("w")}/>')
            sect_pr.append(cols)
        
        if i == 0:
            # First section: single column for title/author block
            cols.set(qn('w:num'), '1')
        else:
            # Body sections: two columns
            # w:space expects value in TWIPS (1/1440 of an inch).
            # 0.25 inches = 360 twips.
            cols.set(qn('w:num'), '2')
            cols.set(qn('w:space'), '360')
            cols.set(qn('w:equalWidth'), '1')
    
    changes.append("Page: 8.5x11in, margins: T=0.75in B=1.0in L=R=0.625in")
    changes.append("Two-column layout for body sections")
    
    # ─── 2. Title paragraph ──────────────────────────────────────────
    print("  [2/7] Formatting title...")
    for i, p in enumerate(doc.paragraphs):
        if is_title_paragraph(i, p.text):
            set_paragraph_format(p,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=Pt(0),
                space_after=Pt(12),
                line_spacing=None,
            )
            set_all_runs_font(p, size=IEEE_TITLE_SIZE, bold=True, italic=False, small_caps=False)
            changes.append(f"Title: 24pt bold centered")
            break
    
    # ─── 3. Author/affiliation tables ────────────────────────────────
    print("  [3/7] Formatting author blocks...")
    for table in doc.tables:
        if is_author_table(table):
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                           space_before=Pt(0), space_after=Pt(0))
                        for run in p.runs:
                            text = run.text.strip()
                            # Author names: 11pt, regular
                            # Check if it's a name (first row) or affiliation
                            if any(name in text for name in ['Pooja', 'Sivaprakash', 'Karthikeyini', 'Nishanth', 'Nithin']):
                                set_run_font(run, size=IEEE_AUTHOR_SIZE, bold=False, italic=False)
                            elif '@' in text or 'mail' in text.lower():
                                # Email: 10pt, not italic
                                set_run_font(run, size=IEEE_AFFILIATION_SIZE, bold=False, italic=False)
                            elif text:
                                # Affiliations: 10pt italic
                                set_run_font(run, size=IEEE_AFFILIATION_SIZE, bold=False, italic=True)
            changes.append("Authors: 11pt centered; Affiliations: 10pt italic centered")
    
    # ─── 4. Format all body paragraphs ───────────────────────────────
    print("  [4/7] Formatting body paragraphs, headings, abstract, references...")
    prev_was_heading = False
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        
        # Skip title (already handled)
        if is_title_paragraph(i, text):
            prev_was_heading = False
            continue
        
        # ── Abstract ──
        if is_abstract_paragraph(text):
            set_paragraph_format(p,
                alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                space_before=Pt(6),
                space_after=Pt(6),
                line_spacing=IEEE_LINE_SPACING,
                line_spacing_rule=WD_LINE_SPACING.EXACTLY,
                first_line_indent=Pt(0),
            )
            for run in p.runs:
                # "Abstract—" should be bold 9pt
                if 'Abstract' in run.text or 'ABSTRACT' in run.text:
                    set_run_font(run, size=IEEE_ABSTRACT_SIZE, bold=True, italic=False)
                else:
                    set_run_font(run, size=IEEE_ABSTRACT_SIZE, bold=False, italic=False)
            changes.append("Abstract: 9pt, bold label, justified")
            prev_was_heading = False
            continue
        
        # ── Keywords ──
        if is_keywords_paragraph(text):
            set_paragraph_format(p,
                alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                space_before=Pt(3),
                space_after=Pt(12),
                line_spacing=IEEE_LINE_SPACING,
                line_spacing_rule=WD_LINE_SPACING.EXACTLY,
                first_line_indent=Pt(0),
            )
            for run in p.runs:
                if 'Keywords' in run.text or 'Index Terms' in run.text:
                    set_run_font(run, size=IEEE_ABSTRACT_SIZE, bold=True, italic=True)
                else:
                    set_run_font(run, size=IEEE_ABSTRACT_SIZE, bold=False, italic=True)
            changes.append("Keywords: 9pt, bold-italic label, italic text")
            prev_was_heading = False
            continue
        
        # ── Section headings (I. INTRODUCTION, II. RELATED WORK...) ──
        if is_section_heading(text):
            set_paragraph_format(p,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=Pt(12),
                space_after=Pt(6),
                line_spacing=IEEE_LINE_SPACING,
                line_spacing_rule=WD_LINE_SPACING.EXACTLY,
                first_line_indent=Pt(0),
            )
            set_all_runs_font(p, size=IEEE_SECTION_HEADING_SIZE, bold=False,
                            italic=False, small_caps=True)
            prev_was_heading = True
            continue
        
        # ── Subsection headings (1.1, 2.3, A. etc.) ──
        if is_subsection_heading(text):
            set_paragraph_format(p,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=Pt(6),
                space_after=Pt(3),
                line_spacing=IEEE_LINE_SPACING,
                line_spacing_rule=WD_LINE_SPACING.EXACTLY,
                first_line_indent=Pt(0),
            )
            set_all_runs_font(p, size=IEEE_SUBSECTION_HEADING_SIZE, bold=False,
                            italic=True, small_caps=False)
            prev_was_heading = True
            continue
        
        # ── Sub-subsection headings (check for bold + short text that look like headings) ──
        # These are paragraphs that are bold, not too long, and look like headings
        if p.runs and all(r.font.bold for r in p.runs if r.text.strip()):
            if len(text) < 80 and not is_reference(text) and i > 5:
                # Check if it matches known heading patterns  
                heading_patterns = [
                    r'^\d+\.\d+\.\d+\s',  # 2.7.1, 3.4.2 etc
                    r'^Level \d',           # Level 1, Level 2
                    r'^Audio ',             # Audio Feature Extraction
                    r'^Text ',              # Text Feature
                    r'^Video ',             # Video Feature
                    r'^Facial ',            # Facial Feature
                    r'^Tabular ',           # Tabular Feature
                ]
                if any(re.match(pat, text) for pat in heading_patterns):
                    set_paragraph_format(p,
                        alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        space_before=Pt(6),
                        space_after=Pt(3),
                        line_spacing=IEEE_LINE_SPACING,
                        line_spacing_rule=WD_LINE_SPACING.EXACTLY,
                        first_line_indent=Pt(0),
                    )
                    set_all_runs_font(p, size=IEEE_SUBSECTION_HEADING_SIZE, bold=False,
                                    italic=True, small_caps=False)
                    prev_was_heading = True
                    continue
                # If it's a section-like heading (all caps or Roman numeral format)
                if text == text.upper() or re.match(r'^[IVX]+\.', text):
                    set_paragraph_format(p,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before=Pt(12),
                        space_after=Pt(6),
                        line_spacing=IEEE_LINE_SPACING,
                        line_spacing_rule=WD_LINE_SPACING.EXACTLY,
                        first_line_indent=Pt(0),
                    )
                    set_all_runs_font(p, size=IEEE_SECTION_HEADING_SIZE, bold=False,
                                    italic=False, small_caps=True)
                    prev_was_heading = True
                    continue
        
        # ── References ──
        if is_reference(text):
            set_paragraph_format(p,
                alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                space_before=Pt(0),
                space_after=Pt(1),
                line_spacing=IEEE_LINE_SPACING,
                line_spacing_rule=WD_LINE_SPACING.EXACTLY,
                first_line_indent=Pt(0),
            )
            # Apply 8pt to all runs; if no explicit runs, set on paragraph
            if p.runs:
                set_all_runs_font(p, size=IEEE_REFERENCE_SIZE, bold=False, italic=False)
            else:
                # Create a run if there are none
                run = p.add_run()
                run.text = ""
                set_run_font(run, size=IEEE_REFERENCE_SIZE)
            
            # Set hanging indent for references: 0.25in left indent, -0.25in first line
            pf = p.paragraph_format
            pf.left_indent = Inches(0.25)
            pf.first_line_indent = Inches(-0.25)
            
            prev_was_heading = False
            continue
        
        # ── Regular body text ──
        indent = Pt(0) if prev_was_heading else IEEE_FIRST_LINE_INDENT
        set_paragraph_format(p,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            space_before=IEEE_PARA_SPACING_BEFORE,
            space_after=IEEE_PARA_SPACING_AFTER,
            line_spacing=IEEE_LINE_SPACING,
            line_spacing_rule=WD_LINE_SPACING.EXACTLY,
            first_line_indent=indent,
        )
        set_all_runs_font(p, size=IEEE_BODY_SIZE, bold=False, italic=False, small_caps=False)
        prev_was_heading = False
    
    changes.append("Body: 10pt Times New Roman, justified, 12pt line spacing")
    changes.append("Section heads: 10pt small-caps centered")
    changes.append("Subsection heads: 10pt italic left-aligned")
    changes.append("References: 8pt with hanging indent")
    
    # ─── 5. Format tables ────────────────────────────────────────────
    print("  [5/7] Formatting tables...")
    for ti, table in enumerate(doc.tables):
        if is_author_table(table):
            continue  # Already handled
        
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for p in cell.paragraphs:
                    set_paragraph_format(p,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before=Pt(1),
                        space_after=Pt(1),
                        line_spacing=Pt(10),
                        line_spacing_rule=WD_LINE_SPACING.EXACTLY,
                        first_line_indent=Pt(0),
                    )
                    for run in p.runs:
                        if ri == 0:
                            # Header row: bold 8pt
                            set_run_font(run, size=IEEE_TABLE_SIZE, bold=True, italic=False)
                        else:
                            # Data rows: 8pt regular
                            set_run_font(run, size=IEEE_TABLE_SIZE, bold=False, italic=False)
    
    changes.append("Tables: 8pt, centered, header row bold")
    
    # ─── 6. Set default paragraph style ──────────────────────────────
    print("  [6/7] Setting default document styles...")
    # Update 'Normal' style
    normal_style = doc.styles['Normal']
    normal_style.font.name = IEEE_FONT
    normal_style.font.size = IEEE_BODY_SIZE
    normal_style.font.bold = False
    normal_style.font.italic = False
    normal_style.font.color.rgb = RGBColor(0, 0, 0)
    
    # Set East Asian font
    rPr = normal_style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{IEEE_FONT}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), IEEE_FONT)
    
    # Set paragraph format for Normal style
    pf = normal_style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = IEEE_LINE_SPACING
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    changes.append("Default 'Normal' style: 10pt Times New Roman, justified, 12pt line spacing")
    
    # ─── 7. Remove extra spacing ─────────────────────────────────────
    print("  [7/7] Cleaning up excess whitespace...")
    # Remove any auto-spacing that Word may add
    for p in doc.paragraphs:
        pPr = p._element.find(qn('w:pPr'))
        if pPr is not None:
            # Remove contextualSpacing if present
            ctx = pPr.find(qn('w:contextualSpacing'))
            if ctx is not None:
                pPr.remove(ctx)
    
    changes.append("Removed contextual spacing overrides")
    
    # ─── Save ────────────────────────────────────────────────────────
    doc.save(doc_path)
    
    print(f"\n{'='*70}")
    print(f"  IEEE FORMATTING APPLIED SUCCESSFULLY")
    print(f"{'='*70}")
    print(f"\n  Changes applied:")
    for c in changes:
        print(f"    ✅ {c}")
    print(f"\n  Saved to: {doc_path}")
    print(f"{'='*70}\n")


# ─────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    DOC_PATH = os.path.join(BASE_DIR, "docs", "ml_pipeline", "H5_OMNIFUSION_JOURNAL_PAPER_V4.docx")
    
    if not os.path.exists(DOC_PATH):
        print(f"ERROR: Document not found: {DOC_PATH}")
        print("Run update_paper.py first to generate V4.")
        exit(1)
    
    format_ieee(DOC_PATH)
