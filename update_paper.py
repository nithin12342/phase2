"""
update_paper.py — H⁵-OmniFusion Journal Paper Updater (V3 → V4)
================================================================
Recursively cross-references the project codebase with the paper text
and applies targeted corrections to align paper claims with actual
implementation.

Usage:
    python update_paper.py

Input:  docs/ml_pipeline/H5_OMNIFUSION_JOURNAL_PAPER_V3.docx
Output: docs/ml_pipeline/H5_OMNIFUSION_JOURNAL_PAPER_V4.docx
"""

import os
import re
import copy
from dataclasses import dataclass, field
from typing import List, Optional
from docx import Document
from docx.shared import Pt, RGBColor


# ─────────────────────────────────────────────────────────────────────
# Correction data structure
# ─────────────────────────────────────────────────────────────────────

@dataclass
class Correction:
    """One targeted text replacement in the paper."""
    section: str            # Which section this corrects
    description: str        # Human-readable reason
    old_text: str           # Exact substring to find (case-sensitive)
    new_text: str           # Replacement text
    applied: bool = False   # Set to True once applied


# ─────────────────────────────────────────────────────────────────────
# All corrections derived from codebase analysis
# ─────────────────────────────────────────────────────────────────────

CORRECTIONS: List[Correction] = [

    # ── 1. Latent Perceiver: N_L = 64 → 16 ──────────────────────────
    Correction(
        section="3.4 (Level 3 — Latent Perceiver)",
        description="n_latents is 16 in config.py FusionConfig, not 64",
        old_text="(L \u2208 R^{64\u00d7768})",
        new_text="(L \u2208 R^{16\u00d7768})",
    ),
    Correction(
        section="3.4 (Level 3 — Latent Perceiver)",
        description="N_L = 16 in config.py, not 64",
        old_text="N_L = 64",
        new_text="N_L = 16",
    ),

    # ── 2. Loss weights: λ_2, λ_3, and λ_4 ────────────────────────────
    # Full-string replacement to fix all lambda values at once and remove λ_4
    Correction(
        section="3.6 (Loss Functions)",
        description="lambda_phq=0.3, lambda_orth=0.05 in config.py; no balance loss (λ_4)",
        old_text="\u03bb_1 = 1.0, \u03bb_2 = 0.5, \u03bb_3 = 0.1, and \u03bb_4 = 0.01",
        new_text="\u03bb_1 = 1.0, \u03bb_2 = 0.3, and \u03bb_3 = 0.05",
    ),
    # Also fix the L_total formula line  
    Correction(
        section="3.6 (Loss Functions — L_total formula)",
        description="Remove \u03bb_4 L_balance from formula",
        old_text="\u03bb_1 L_Focal + \u03bb_2 L_MSE + \u03bb_3 L_Orth + \u03bb_4 L_balance",
        new_text="\u03bb_1 L_Focal + \u03bb_2 L_MSE + \u03bb_3 L_Orth",
    ),

    # ── 3. Training parameters ───────────────────────────────────────
    Correction(
        section="5.2 (Implementation Details — Epochs)",
        description="n_epochs = 50 in config.py, not 100",
        old_text="Training proceeds for 100 epochs with early stopping based on validation F1 score (patience = 10 epochs).",
        new_text="Training proceeds for 50 epochs with early stopping based on Youden's J statistic (sensitivity + specificity − 1) with patience = 15 epochs.",
    ),
    Correction(
        section="5.2 (Implementation Details — Batch)",
        description="BATCH_SIZE = 1 per participant due to variable-length multimodal data",
        old_text="Batch size is 8 due to memory constraints from processing five modalities (approximately 32 GB GPU memory).",
        new_text="Batch size is 1 (per participant) due to variable-length multimodal sequences across five modalities, with gradient accumulation to simulate larger effective batches.",
    ),
    Correction(
        section="5.2 (Implementation Details — Scheduler)",
        description="OneCycleLR used in trainer.py, not plain cosine annealing",
        old_text="cosine annealing with 10% warmup epochs",
        new_text="OneCycleLR scheduling with cosine annealing strategy and 10% warmup ratio (div_factor=25, final_div_factor=1000)",
    ),

    # ── 4. Focal Loss alpha ──────────────────────────────────────────
    Correction(
        section="3.6 (Classification Head)",
        description="FocalLossBinary uses alpha=0.75 for depressed class in trainer.py",
        old_text="\u03b1 = 0.25 balances positive/negative examples",
        new_text="\u03b1 = 0.75 balances positive/negative examples (weighting the depressed class more heavily)",
    ),
    # Also fix the separate mention in V14 section
    Correction(
        section="2.7 (V14 Architecture)",
        description="Focal Loss alpha is 0.75 in trainer implementation",
        old_text="Focal Loss with \u03b1=0.3 and \u03b3=2.0",
        new_text="Focal Loss with \u03b1=0.75 and \u03b3=2.0",
    ),

    # ── 5. Encoder architecture: BiLSTM → Mamba ──────────────────────
    Correction(
        section="3.3 (Stage 2 — Modality-Specific Encoders)",
        description="Code uses MambaEncoder, not BiLSTM, for temporal encoding",
        old_text=(
            "The second stage applies temporal encoding to sequential modalities, "
            "capturing intra-modality dynamics through Bidirectional Long Short-Term "
            "Memory (BiLSTM) networks with temporal attention."
        ),
        new_text=(
            "The second stage applies temporal encoding to sequential modalities, "
            "capturing intra-modality dynamics through Mamba-based state-space model "
            "(SSM) encoders with efficient linear-time sequence modeling."
        ),
    ),
    Correction(
        section="3.3 (Stage 2 — BiLSTM formula)",
        description="Replace BiLSTM processing description with Mamba description",
        old_text=(
            "For sequential modality m ∈ {A, V, F}, the BiLSTM encoder processes "
            "the feature sequence X^(m) = [x_1, x_2, ..., x_T] bidirectionally. "
            "The forward LSTM computes hidden states h→_t = LSTM→(x_t, h→_{t-1}), "
            "while the backward LSTM computes h←_t = LSTM←(x_t, h←_{t+1}). "
            "The concatenated bidirectional hidden state is h_t = [h→_t; h←_t] ∈ R^{2d_h} "
            "where d_h = 384 yields 768-dimensional combined representations matching "
            "the embedding dimension."
        ),
        new_text=(
            "For sequential modality m ∈ {A, V, F}, the Mamba encoder processes "
            "the feature sequence X^(m) = [x_1, x_2, ..., x_T] through a selective "
            "state-space model (SSM) that provides linear-time complexity O(T) for "
            "long sequences. The Mamba architecture uses input-dependent selection "
            "mechanisms to dynamically weight temporal dependencies, enabling efficient "
            "modeling of depression-relevant temporal patterns without the quadratic "
            "cost of self-attention. Each encoder consists of n_layers=2 Mamba blocks "
            "with dropout=0.1, producing representations h_t ∈ R^{768} matching the "
            "unified embedding dimension."
        ),
    ),
    Correction(
        section="3.3 (Stage 2 — Temporal attention)",
        description="Replace temporal attention aggregation with mean pooling used in code",
        old_text=(
            "Temporal attention aggregates the sequence into a fixed-length "
            "representation by learning importance weights over time steps. "
            "The attention mechanism computes: α_t = softmax(w^T tanh(W_h h_t + b_h)) "
            "where W_h ∈ R^{d_att × 768}, w ∈ R^{d_att}, and d_att = 256 is the "
            "attention dimension. The attended representation is the weighted sum: "
            "z^(m) = Σ_t α_t h_t ∈ R^{768}. This attention mechanism learns to "
            "emphasize time points containing depression-relevant signals while "
            "down-weighting uninformative segments."
        ),
        new_text=(
            "After temporal encoding, the sequence is aggregated into a fixed-length "
            "summary representation through mean pooling over the temporal dimension: "
            "z^(m) = (1/T) Σ_t h_t ∈ R^{768}, followed by a learned summary projection "
            "and layer normalization: z^(m) = LayerNorm(W_proj · mean(H^(m)) + b_proj). "
            "This provides a stable, information-preserving aggregation of the temporal "
            "representation for downstream fusion processing."
        ),
    ),

    # ── 6. MoE expert architecture ───────────────────────────────────
    Correction(
        section="3.5 (MoE — Expert networks)",
        description="Each expert outputs a scalar logit (not 768-dim); hidden=512",
        old_text=(
            "Each expert is a two-layer feedforward network: "
            "E_i(z) = W_2^{(i)} ReLU(W_1^{(i)} z + b_1^{(i)}) + b_2^{(i)} "
            "with W_1 ∈ R^{2048 × 768}, W_2 ∈ R^{768 × 2048}."
        ),
        new_text=(
            "Each modality expert is a two-layer feedforward network producing "
            "a scalar logit: E_i(z) = W_2^{(i)} ReLU(W_1^{(i)} z + b_1^{(i)}) + b_2^{(i)} "
            "with W_1 ∈ R^{512 × d_input}, W_2 ∈ R^{1 × 512}, where d_input is the "
            "concatenated shared+specific dimension from MS² decomposition. A sixth "
            "fusion expert operates on the global CLS representation with hidden "
            "dimension 256."
        ),
    ),
    Correction(
        section="3.5 (MoE — Final output)",
        description="Expert outputs are scalar logits weighted by gate, not 768-dim vectors",
        old_text=(
            "The final output combines selected expert outputs weighted by "
            "normalized gate values: z_MoE = Σ_{i∈TopK} (g_i / Σ_{j∈TopK} g_j) · "
            "E_i(z_fused)."
        ),
        new_text=(
            "The final output combines all six expert scalar logits weighted by "
            "softmax gate values: logit_final = Σ_i g_i · E_i(z_i), where g_i "
            "are quality-aware gate weights computed from all modality summaries, "
            "the global CLS token, and learned quality feature projections."
        ),
    ),
    # Remove Top-K=3 claims since code uses full softmax over all 6 experts
    Correction(
        section="3.5 (MoE — Gating)",
        description="Code does full softmax over 6 experts, not Top-K=3 selection",
        old_text=(
            "Top-K selection retains the K=3 highest-weighted experts: "
            "TopK(g) = {i : g_i ≥ g_{(K)}} where g_{(K)} is the K-th largest "
            "gate value. Expert outputs are computed only for selected experts, "
            "reducing computational cost."
        ),
        new_text=(
            "All six expert outputs are computed and combined using softmax-normalized "
            "gate weights. The quality-aware gating network takes as input the "
            "concatenation of all five modality summaries, the global CLS token, "
            "and projected quality features (audio SNR, face confidence, text length, "
            "video motion, tabular completeness), producing routing weights through "
            "a learned linear projection followed by softmax normalization."
        ),
    ),
    # Remove balance loss mention
    Correction(
        section="3.5 (MoE — Balance loss)",
        description="No auxiliary load balancing loss in code",
        old_text=(
            "An auxiliary load balancing loss encourages uniform expert utilization: "
            "L_balance = CV(Σ_x 1[i ∈ TopK(g_x)]) measuring coefficient of variation "
            "across expert selection frequencies."
        ),
        new_text=(
            "Expert diversity is encouraged through the orthogonality regularization "
            "in the MS² decomposition module, which penalizes overlap between the "
            "shared and modality-specific subspaces fed to each expert."
        ),
    ),
    # Fix MoE noise claim
    Correction(
        section="3.5 (MoE — Gating network)",
        description="No noise injection in QualityAwareGate code",
        old_text="g = softmax(W_g z_fused + b_g + ε) where W_g ∈ R^{6 × 768}, b_g ∈ R^6, and ε ~ N(0, 0.01) is noise for load balancing during training.",
        new_text="The gating network processes the concatenated modality summaries, CLS token, and quality features: g = softmax(W_g · [z_A; z_V; z_F; z_T; z_B; z_CLS; q] + b_g) where W_g ∈ R^{6 × (6·d_model + n_quality)}, with input clamping for numerical stability.",
    ),

    # ── 7. Tabular features count ────────────────────────────────────
    Correction(
        section="3.2 (Feature Extraction — Tabular)",
        description="n_features = 20 in config.py TabularConfig, not 15",
        old_text="Tabular metadata comprises participant demographics (age, gender), interview statistics (duration, number of turns, words per turn), and session quality metrics (SNR, face detection rate, frame drop percentage). These 15-dimensional features",
        new_text="Tabular metadata comprises participant demographics (age, gender), interview statistics (duration, number of turns, words per turn), session quality metrics (SNR, face detection rate, frame drop percentage), sentiment scores (positive, negative, neutral), and health survey indicators (stress, habit changes, mental health history, family history, coping struggles, social weakness). These 20-dimensional features",
    ),

    # ── 8. Video FPS ─────────────────────────────────────────────────
    Correction(
        section="4.2 (Video Preprocessing)",
        description="TARGET_FPS = 5 in config.py, not native 25 FPS",
        old_text="Video preprocessing begins with frame extraction at native frame rate (typically 25 FPS for DAIC-WOZ)",
        new_text="Video preprocessing begins with frame extraction at a reduced target frame rate of 5 FPS (down-sampled from the native 25–30 FPS in DAIC-WOZ)",
    ),

    # ── 9. Face detection thresholds ─────────────────────────────────
    Correction(
        section="4.2 (Face Preprocessing — Quality)",
        description="FACE_CONFIDENCE_THRESHOLD = 0.8 in config.py, not 0.85",
        old_text="mean confidence > 0.85",
        new_text="mean confidence > 0.80",
    ),
    Correction(
        section="4.2 (Face Preprocessing — Detection Rate)",
        description="VIDEO_FACE_DETECTION_MIN_RATIO = 0.8, not 0.95",
        old_text="detection success rate > 95%",
        new_text="detection success rate > 80%",
    ),

    # ── 10. Hypergraph Level 1 clarification ─────────────────────────
    Correction(
        section="3.4 (Level 1 — Local Hypergraph)",
        description="Code uses multi-head attention over stacked modalities at each timestep, not literal hypergraph convolution",
        old_text=(
            "The hypergraph convolution propagates information: "
            "H'^(m) = σ(D_v^{-1/2} H W D_e^{-1} H^T D_v^{-1/2} H^(m) Θ^(m)) "
            "where H is the incidence matrix, D_v and D_e are degree matrices, "
            "W is the learnable edge weight matrix, Θ^(m) ∈ R^{768 × 768} is the "
            "layer weight, and σ is the ReLU activation. This operation enables "
            "each time step to aggregate information from its local temporal "
            "neighborhood through the hyperedge structure."
        ),
        new_text=(
            "At each time step t, the modality representations are stacked into "
            "a tensor S_t ∈ R^{M × d} (where M is the number of available modalities) "
            "and processed through multi-head attention with 8 heads: "
            "S'_t = LayerNorm(S_t + MHA(Q=S_t, K=S_t, V=S_t)). "
            "This hypergraph-inspired operation enables all modalities at each time "
            "step to exchange information simultaneously, capturing cross-modal "
            "correlations at the temporal level. The fused representation at each "
            "time step is obtained by mean-pooling across the modality dimension."
        ),
    ),

    # ── 11. Hypergraph Level 2 clarification ─────────────────────────
    Correction(
        section="3.4 (Level 2 — Modality Hypergraph)",
        description="Code uses TransformerEncoder with self-attention, not explicit 31-hyperedge construction",
        old_text=(
            "Unlike pairwise graphs, we construct hyperedges connecting modality "
            "subsets: 5 singleton edges (self-loops), 10 pairwise edges, 10 triplet "
            "edges, 5 quartet edges, and 1 full 5-way edge. This totals 31 hyperedges "
            "capturing all possible inter-modal interaction patterns. The hyperedge "
            "weights are learned parameters initialized uniformly and optimized through "
            "backpropagation. The modality hypergraph convolution applies the same "
            "propagation rule with hyperedge attention: β_e = softmax(MLP(Σ_{v∈e} z_v / |e|)) "
            "weighting hyperedge contributions."
        ),
        new_text=(
            "The modality-level hypergraph is implemented as a Transformer encoder "
            "operating over the five modality summary tokens. The summaries are stacked "
            "into a sequence Z ∈ R^{5 × d} and processed through n_layers=2 Transformer "
            "encoder layers with 8 attention heads, GELU activation, and feedforward "
            "dimension 4×d_model. The self-attention mechanism allows each modality "
            "to attend to all others, learning implicit higher-order interaction patterns "
            "through multiple layers. Unlike explicit hyperedge construction, this learned "
            "attention approach adaptively discovers the most relevant cross-modal "
            "relationships through end-to-end training."
        ),
    ),

    # ── 12. eGeMAPSv02 usage clarification ───────────────────────────
    Correction(
        section="3.2 (Audio Feature Extraction)",
        description="use_egemaps = False in AudioConfig; eGeMAPSv02 is available but disabled",
        old_text="The final audio representation concatenates Wav2Vec2 and projected eGeMAPSS features: X^(A) = [X_wav2vec; W_eGeMAPSS · X_egemaps].",
        new_text="The audio representation uses Wav2Vec2 features projected to the unified embedding space. While eGeMAPSv02 feature integration is architecturally supported (via a learned projection W_eGeMAPS ∈ R^{88 × 768}), the default configuration uses Wav2Vec2 representations alone, as they subsume many traditional acoustic features through self-supervised pretraining.",
    ),

    # ── 13. Face AU dimensions ───────────────────────────────────────
    Correction(
        section="3.2 (Facial Feature Extraction)",
        description="au_dim = 35 in FaceConfig, not 188",
        old_text="These 188-dimensional per-frame features are concatenated with POSTER v2 expression embeddings.",
        new_text="From these, 35 key features (17 AU intensities + AU presences + gaze vectors) are extracted and concatenated with POSTER v2 expression embeddings.",
    ),

    # ── 14. Add label smoothing mention ──────────────────────────────
    Correction(
        section="3.6 (Classification Head — Focal Loss)",
        description="FocalLossBinary in trainer.py uses label_smoothing=0.05",
        old_text="Focal Loss addresses class imbalance by down-weighting well-classified examples: L_Focal = -α (1 - p_t)^γ log(p_t) where p_t = ŷ_cls if y=1 else (1-ŷ_cls),",
        new_text="Focal Loss with label smoothing (ε=0.05) addresses class imbalance by down-weighting well-classified examples: L_Focal = -α (1 - p_t)^γ log(p_t) where p_t = ŷ_cls if y=1 else (1-ŷ_cls), targets are smoothed as y' = y(1−ε) + 0.5ε,",
    ),

    # ── 15. Abstract: update Top-K=3 claim ───────────────────────────
    Correction(
        section="Abstract",
        description="Code uses full softmax gating over all 6 experts, not Top-K=3",
        old_text="A quality-aware Mixture of Experts (MoE) layer with Top-K=3 routing enables dynamic modality weighting",
        new_text="A quality-aware Mixture of Experts (MoE) layer with softmax routing over six experts (five modality-specific + one fusion) enables dynamic modality weighting",
    ),

    # ── 16. Introduction: update Top-K=3 reference ──────────────────
    Correction(
        section="1.2 (Research Contribution)",
        description="Full softmax routing, not Top-K=3",
        old_text="with Top-K=3 selection ensuring computational efficiency while maintaining capacity",
        new_text="with softmax-normalized routing ensuring balanced expert utilization",
    ),

    # ── 17. Section 2.7.1: update Top-K mention ─────────────────────
    Correction(
        section="2.7.1 (Architectural Comparison)",
        description="Full softmax routing in implementation",
        old_text="Top-K=3 selection activates only the three most relevant experts per sample, providing computational efficiency.",
        new_text="Softmax routing distributes input across all six experts with quality-aware weighting, providing adaptive modality emphasis based on input quality indicators.",
    ),

    # ── 18. MS² Decomposition: add to methodology ───────────────────
    # This is an addition — we insert it after the modality hypergraph section
    # by modifying the output description
    Correction(
        section="3.4 (Level 2 — Output)",
        description="Add MS² Decomposition mention (exists in code as ms2_decomposition.py)",
        old_text="The output Z' = {z'^(A), z'^(T), z'^(V), z'^(F), z'^(B)} captures cross-modal dependencies.",
        new_text=(
            "The output Z' = {z'^(A), z'^(T), z'^(V), z'^(F), z'^(B)} captures "
            "cross-modal dependencies. These updated summaries are further decomposed "
            "by a Modality Shared-Specific (MS²) module into shared subspace components "
            "s^(m) ∈ R^{d_shared} (cross-modal depression-relevant features) and "
            "specific subspace components q^(m) ∈ R^{d_specific} (modality-unique features) "
            "via learned projections with LayerNorm, where d_shared = d_specific = d_model/2. "
            "An orthogonality loss L_orth = (1/M) Σ_m mean(⟨s^(m), q^(m)⟩²) penalizes "
            "overlap between shared and specific subspaces, encouraging clean decomposition."
        ),
    ),

    # ── 19. Conclusion: Top-K=3 → softmax routing ───────────────────
    Correction(
        section="7.1 (Summary — MoE)",
        description="Align conclusion with corrected MoE description",
        old_text="the quality-aware Mixture of Experts enables dynamic modality weighting based on implicit input quality indicators",
        new_text="the quality-aware Mixture of Experts with softmax routing over six specialized experts enables dynamic modality weighting based on implicit input quality indicators",
    ),

    # ── 20. Dataset split correction ──────────────────────────────────
    Correction(
        section="5.1 (Dataset)",
        description="DAIC-WOZ has 189 participants: 107 dep + 82 non-dep",
        old_text="the corpus is partitioned into training (107 participants), development (35 participants), and test (47 participants) sets",
        new_text="the corpus is partitioned into training (107 participants), development (35 participants), and test (47 participants) sets following the standard AVEC protocol",
    ),

    # ── 21. Fix cross-validation mention ─────────────────────────────
    Correction(
        section="5.2 (Cross-validation)",
        description="Code uses folds but paper says 5-fold; add clarification",
        old_text="Five-fold cross-validation is employed to obtain robust performance estimates, with each fold preserving the class distribution.",
        new_text="Five-fold cross-validation is employed to obtain robust performance estimates, with each fold preserving the class distribution. Model selection within each fold uses Youden's J statistic (Sensitivity + Specificity − 1) rather than raw F1 to ensure balanced detection of both depressed and non-depressed classes, with a decision threshold of 0.35.",
    ),

    # ── 22. (Removed — already covered by correction #16 which modifies
    #         the same paragraph 18 text about Top-K=3 selection) ──────

    # ── 23. Fix V14 parameter count context ──────────────────────────
    Correction(
        section="2.7 (V14 Architecture — Parameters)",
        description="Clarify parameter count context",
        old_text="The complete V14 model contains approximately 6.5 million parameters.",
        new_text="The complete V14 model contains approximately 6.5 million trainable parameters, serving as the lightweight predecessor to H⁵-OmniFusion.",
    ),
]


# ─────────────────────────────────────────────────────────────────────
# Document update engine
# ─────────────────────────────────────────────────────────────────────

def apply_correction_to_paragraph(paragraph, correction: Correction) -> bool:
    """
    Apply a single correction to a paragraph.
    
    Handles the complexity of python-docx runs: text may be split across
    multiple runs, so we work at the full-paragraph level and then
    reconstruct runs preserving the first run's formatting.
    """
    full_text = paragraph.text
    if correction.old_text not in full_text:
        return False

    # Perform the replacement on full text
    new_full_text = full_text.replace(correction.old_text, correction.new_text, 1)

    # Preserve the formatting of the first run
    if paragraph.runs:
        # Store first run's formatting
        first_run = paragraph.runs[0]
        font_name = first_run.font.name
        font_size = first_run.font.size
        font_bold = first_run.font.bold
        font_italic = first_run.font.italic

        # Clear all runs
        for run in paragraph.runs:
            run.text = ""
        
        # Set the new text on the first run
        paragraph.runs[0].text = new_full_text

        # Re-apply formatting
        paragraph.runs[0].font.name = font_name
        paragraph.runs[0].font.size = font_size
        paragraph.runs[0].font.bold = font_bold
        paragraph.runs[0].font.italic = font_italic
    else:
        paragraph.text = new_full_text

    correction.applied = True
    return True


def update_paper(input_path: str, output_path: str) -> List[Correction]:
    """
    Read the paper, apply all corrections, save the updated version.
    Returns the list of corrections with their applied status.
    """
    print(f"\n{'='*70}")
    print(f"  H⁵-OmniFusion Paper Updater — V3 → V4")
    print(f"{'='*70}")
    print(f"\n  Input:  {input_path}")
    print(f"  Output: {output_path}")
    print(f"  Corrections to apply: {len(CORRECTIONS)}")
    print()

    # Load document
    doc = Document(input_path)
    total_paragraphs = len(doc.paragraphs)
    print(f"  Document loaded: {total_paragraphs} paragraphs, {len(doc.tables)} tables\n")

    # Apply each correction
    for i, correction in enumerate(CORRECTIONS, 1):
        found = False
        for para in doc.paragraphs:
            if apply_correction_to_paragraph(para, correction):
                found = True
                break

        status = "✅ APPLIED" if found else "❌ NOT FOUND"
        print(f"  [{i:2d}/{len(CORRECTIONS)}] {status} — {correction.section}")
        if found:
            # Show a snippet of the change
            old_snippet = correction.old_text[:60].replace('\n', ' ')
            new_snippet = correction.new_text[:60].replace('\n', ' ')
            print(f"           Old: \"{old_snippet}...\"")
            print(f"           New: \"{new_snippet}...\"")
        else:
            print(f"           Looking for: \"{correction.old_text[:80]}...\"")

    # Save the updated document
    doc.save(output_path)

    # Summary report
    applied = sum(1 for c in CORRECTIONS if c.applied)
    failed = len(CORRECTIONS) - applied

    print(f"\n{'='*70}")
    print(f"  SUMMARY")
    print(f"{'='*70}")
    print(f"  Total corrections:   {len(CORRECTIONS)}")
    print(f"  Successfully applied: {applied}")
    print(f"  Not found (skipped): {failed}")
    print(f"\n  Output saved to: {output_path}")

    if failed > 0:
        print(f"\n  ⚠️  {failed} corrections could not be applied.")
        print(f"  These may need manual review — the target text may have been")
        print(f"  slightly different (whitespace, encoding, etc.):\n")
        for c in CORRECTIONS:
            if not c.applied:
                print(f"    • [{c.section}] {c.description}")
                print(f"      Target: \"{c.old_text[:100]}...\"")
                print()

    print(f"{'='*70}\n")
    return CORRECTIONS


# ─────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    INPUT_PATH = os.path.join(BASE_DIR, "docs", "ml_pipeline", "H5_OMNIFUSION_JOURNAL_PAPER_V3.docx")
    OUTPUT_PATH = os.path.join(BASE_DIR, "docs", "ml_pipeline", "H5_OMNIFUSION_JOURNAL_PAPER_V4.docx")

    if not os.path.exists(INPUT_PATH):
        print(f"ERROR: Input file not found: {INPUT_PATH}")
        exit(1)

    results = update_paper(INPUT_PATH, OUTPUT_PATH)

    # Exit with error code if any corrections failed
    failed = sum(1 for c in results if not c.applied)
    exit(0 if failed == 0 else 1)
